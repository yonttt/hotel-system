import docx
from docx.shared import RGBColor

doc = docx.Document('c:/Users/yonat/hotel-system/Form4-Implementation.docx')

markers = []

for i, p in enumerate(doc.paragraphs):
    text = p.text
    if "Axios Interceptors" in text:
        markers.append((i+1, "👉 [INSERT SCREENSHOT OF FRONTEND CODE HERE: hotel-react-frontend/src/services/api.js]"))
    elif "checkin.py / checkout.py" in text:
        markers.append((i+1, "👉 [INSERT SCREENSHOT OF BACKEND CODE HERE: hotel-python-backend/app/api/hotel_reservations.py]"))
    elif "Pydantic schemas" in text:
        markers.append((i+1, "👉 [INSERT SCREENSHOT OF DATABASE MODELS HERE: hotel-python-backend/app/models/__init__.py]"))
        markers.append((i+1, "👉 [INSERT SCREENSHOT OF DATABASE HERE: show your MySQL tables in phpMyAdmin/Workbench]"))
    elif "Tailwind CSS implements" in text:
        markers.append((i+1, "👉 [INSERT SCREENSHOT OF UI HERE: hotel-react-frontend Admin Dashboard]"))
        markers.append((i+1, "👉 [INSERT SCREENSHOT OF UI HERE: WEBSITE-Hotel Home or Booking Page]"))
    elif "Overall Security Observability" in text:
        markers.append((i+1, "👉 [INSERT SCREENSHOT OF INTEGRATION HERE: Network Tab (F12) or Postman showing API Request]"))
        markers.append((i+1, "👉 [INSERT SCREENSHOT OF INTEGRATION HERE: FastAPI Swagger UI (http://localhost:8000/docs)]"))

# Sort descending to avoid shifting indices when inserting
markers.sort(key=lambda x: x[0], reverse=True)

for idx, marker_text in markers:
    if idx < len(doc.paragraphs):
        new_p = doc.paragraphs[idx].insert_paragraph_before()
    else:
        new_p = doc.paragraphs[-1].insert_paragraph_before()
        
    run = new_p.add_run(marker_text)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0) # Red color

doc.save('c:/Users/yonat/hotel-system/Form4-Implementation.docx')
print("Successfully marked the document.")